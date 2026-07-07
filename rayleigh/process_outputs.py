"""`rayleigh process_outputs` — reduce cell data into an experiment's preregistered
outputs, findings, and the datestamped .docx write-up.

rayleigh does NOT compute the analytical products itself. For each experiment it re-expands its
cells (the same expansion conduct_exp ran), loads each cell's output into a tidy table, and writes
that table plus a **base R + ggplot2** script (see rayleigh.r_analysis) that produces every figure,
table, summary statistic, and regression. `process_outputs` runs the script and then ASSEMBLES the
report around what R produced. The script + data are durable: `Rscript results/analysis/<E>.R`
regenerates every analytical product at any time, independently of rayleigh. The finding is stated
honestly — the R-computed observed summary next to the preregistered `expected_direction`, honoring
`mode: confirmatory | exploratory`. No auto-verdict gate; honesty lives in the wording.

Reading a cell's output(s) uses a loader (a cell may write several artifacts):
  - default: each artifact as a JSON dict of scalars, or parquet/csv numeric column means,
    merged (keys prefixed by artifact name when there are several);
  - override: `code.output_adapter.load: "module:callable"` —
    callable(outputs: dict[name, path]) -> {metric: value}.

Deliverables, two audiences:
  - for raconteur (its load_results ingests results/ *.md / *.json / *.csv): `RESULTS.md`
    (prose), `findings.json` (structured per-experiment: prereg + observed finding + artifact
    pointers), and `tables/*.csv` (the R-computed numbers). This is rayleigh's reason to exist.
  - for the human: `results/{cycle}_{project}_results_{ra}.docx` (the ra/DCR review cycle;
    python-docx, degrades to RESULTS.md if unavailable).

Every figure is written by R to `results/figures/` as PNG (embeds in the .docx / markdown preview)
and EPS (vector, for publication); findings.json lists both formats per figure.
"""

import importlib
import json
import sys
from datetime import date, datetime, timezone
from pathlib import Path

import yaml

from rayleigh import __version__
from rayleigh.conduct_exp import expand_cells, resolve_cell_outputs, add_import_paths  # reuse cell logic
from rayleigh.r_analysis import analyze as r_analyze
from rayleigh.spec import active_spec_path


def log(msg: str) -> None:
    print(f"[rayleigh process_outputs] {msg}", flush=True)


# --------------------------------------------------------------------- loading cells
def _flatten(d: dict, prefix: str = "") -> dict:
    out = {}
    for k, v in d.items():
        key = f"{prefix}{k}"
        if isinstance(v, dict):
            out.update(_flatten(v, key + "."))
        else:
            out[key] = v
    return out


def _load_one(path: Path) -> dict:
    """Scalar metrics from one artifact file (JSON dict, or numeric column means of parquet/csv)."""
    if path.suffix == ".json":
        d = json.loads(path.read_text())
        if not isinstance(d, dict):
            return {}
        return {k: v for k, v in _flatten(d).items()
                if isinstance(v, (int, float)) and not isinstance(v, bool)}
    if path.suffix in (".parquet", ".pq", ".csv"):
        import pandas as pd
        df = pd.read_csv(path) if path.suffix == ".csv" else pd.read_parquet(path)
        num = df.select_dtypes("number")
        return {c: float(num[c].mean()) for c in num.columns}
    raise ValueError(f"no default loader for '{path.suffix}' — set code.output_adapter.load")


def _default_load(outputs: dict) -> dict:
    """Merge one cell's scalar metrics across its artifact(s) (outputs = {name: path}).
    With several artifacts, keys are prefixed by name (`summary.rmse`) so they don't collide."""
    multi = len(outputs) > 1
    merged = {}
    for name, path in outputs.items():
        for k, v in _load_one(Path(path)).items():
            merged[f"{name}.{k}" if multi else k] = v
    return merged


def _get_loader(spec: dict, code_dir: Path, results: Path):
    ep = ((spec.get("code") or {}).get("output_adapter") or {}).get("load")
    if not ep:
        return _default_load
    # the loader shim may live in code/ OR results/ (same as the run-adapter) — search both
    add_import_paths(code_dir, results)
    mod_name, _, fn_name = ep.partition(":")
    return getattr(importlib.import_module(mod_name), fn_name)


def _param_keys(exp: dict) -> list:
    """Parameter column names for this experiment (constants + swept/drawn params) — the rest of
    the tidy table's columns are metrics. Covers sweep (axes), conditions, and sobol (categorical
    + continuous); `fixed:` may sit at the experiment or design level."""
    design = exp.get("design") or {}
    keys: list = []

    def add(src):
        if isinstance(src, dict):
            for k in src:
                if k not in keys:
                    keys.append(k)

    add(design.get("fixed"))
    add(exp.get("fixed"))                               # constants are params, not metrics
    kind = str(design.get("kind", "sweep"))
    if kind == "conditions":
        for c in design.get("conditions") or []:
            add(c)
    elif kind == "sobol":
        add(design.get("categorical"))
        add(design.get("continuous"))
    else:                                              # sweep
        add(design.get("axes"))
    return keys


def build_table(exp: dict, adapter: dict, results: Path, loader):
    """Tidy table: one row per cell with output = {**params, seed, **metrics}."""
    import pandas as pd
    eid = str(exp["id"])
    rows, missing = [], 0
    for cell in expand_cells(exp, adapter):
        outputs = resolve_cell_outputs(adapter, results, eid, cell)
        if not all(p.exists() for p in outputs.values()):
            missing += 1
            continue
        try:
            metrics = loader({n: str(p) for n, p in outputs.items()})
        except Exception as e:                       # noqa: BLE001
            log(f"  {eid}: could not load cell {cell['params']} seed={cell['seed']}: "
                f"{type(e).__name__}: {e}")
            metrics = {}
        rows.append({**cell["params"], "seed": cell["seed"], **metrics})
    df = pd.DataFrame(rows)
    param_cols = [k for k in _param_keys(exp) if k in df.columns]
    metric_cols = [c for c in df.columns if c not in param_cols and c != "seed"]
    return df, param_cols, metric_cols, missing


# --------------------------------------------------------------------- finding (from R output)
def _observed_from_summary(exp: dict, summary_df) -> str:
    """Transcribe the primary metric's R-computed range for the finding sentence — a READING of
    R's summary-statistics output, not a Python computation (all stats live in the R script)."""
    if summary_df is None or getattr(summary_df, "empty", True):
        return "data collected; see the figures and the R summary-statistics table"
    primary = (exp.get("metric") or {}).get("name")
    try:
        if primary is not None and primary in summary_df.index:
            row = summary_df.loc[primary]
        else:
            primary, row = str(summary_df.index[0]), summary_df.iloc[0]
        return (f"{primary} ranged {float(row['min']):.3g}–{float(row['max']):.3g} "
                f"(mean {float(row['mean']):.3g}) across the sweep; see the figures")
    except Exception:                                # noqa: BLE001
        return "data collected; see the figures and the R summary-statistics table"


def finding_text(exp: dict, summary_df) -> str:
    """An honest finding: the R-computed observed summary next to the prereg, honoring mode."""
    mode = exp.get("mode", "confirmatory")
    obs = _observed_from_summary(exp, summary_df)
    if mode == "exploratory":
        return f"Exploratory (not preregistered). Observed: {obs}."
    expected = exp.get("expected_direction")
    prefix = f"Expected (preregistered): {expected}. " if expected else ""
    return f"{prefix}Observed: {obs}."




# --------------------------------------------------------------------- assembly
def _fmt(v) -> str:
    return f"{v:.3g}" if isinstance(v, (int, float)) else str(v)


def _pivot_to_md(piv) -> str:
    cols = [str(piv.index.name or "")] + [str(c) for c in piv.columns]
    lines = ["| " + " | ".join(cols) + " |",
             "|" + "|".join("---" for _ in cols) + "|"]
    for idx, row in piv.iterrows():
        lines.append("| " + " | ".join([str(idx)] + [_fmt(v) for v in row]) + " |")
    return "\n".join(lines)


def _build_markdown(project, cycle, brief, items, results: Path) -> str:
    L = [f"# {project} — Results (cycle {cycle})", "", brief, ""]
    for it in items:
        exp = it["exp"]
        L.append(f"## {exp['id']} — {exp.get('title', '')}")
        L.append(f"*Mode: {exp.get('mode', 'confirmatory')}* · "
                 f"{it['n_data']}/{it['n_cells']} cells with data\n")
        if exp.get("question"):
            L.append(f"**Question:** {exp['question']}\n")
        L.append(f"**Finding:** {it['finding']}\n")
        for primary, _all, cap in it["figures"]:
            rel = primary.relative_to(results)   # PNG embeds in the markdown preview
            L.append(f"![{cap}]({rel})\n\n*{cap}*\n")
        for piv, cap in it["tables"]:
            L.append(_pivot_to_md(piv) + f"\n\n*{cap}*\n")
    L.append(f"\n---\n*Generated by rayleigh {__version__}.*")
    return "\n".join(L)


def _build_findings(project, cycle, brief, items, results: Path) -> str:
    """Machine-readable per-experiment findings for the raconteur hand-off.

    rayleigh's reason to exist is to feed raconteur; its load_results() ingests results/*.json
    and *.csv. This is the structured layer (the .docx is the human deliverable) — the
    preregistration, the observed finding, and pointers to the figure/table artifacts.
    """
    doc = {
        "project": project, "cycle": cycle, "brief": brief,
        "rayleigh_version": __version__,
        "generated": datetime.now(timezone.utc).isoformat(),
        "experiments": [],
    }
    for it in items:
        exp = it["exp"]
        metric = exp.get("metric") or {}
        doc["experiments"].append({
            "id": str(exp["id"]),
            "title": exp.get("title", ""),
            "mode": exp.get("mode", "confirmatory"),
            "question": exp.get("question", ""),
            "metric": {"name": metric.get("name"), "reduce": metric.get("reduce")},
            "expected_direction": exp.get("expected_direction", ""),
            "finding": it["finding"],
            "cells": {"with_data": it["n_data"], "total": it["n_cells"]},
            "figures": [{"path": str(primary.relative_to(results)),
                         "formats": {p.suffix.lstrip("."): str(p.relative_to(results))
                                     for p in all_paths},
                         "caption": c}
                        for primary, all_paths, c in it["figures"]],
            "tables": [{"path": str(p.relative_to(results)), "caption": c}
                       for p, c in it.get("table_files", [])],
            "analysis": ({"script": str(Path(a["script"]).relative_to(results)),
                          "data": str(Path(a["data"]).relative_to(results)),
                          "engine": "R (base + ggplot2)", "ran": a.get("ran"), "ok": a.get("ok")}
                         if (a := it.get("analysis")) else None),
        })
    return json.dumps(doc, indent=2)


def _methods_lines(exp: dict) -> list:
    """Human-readable Method bullets: design, fixed params, replications, metric(s)."""
    design = exp.get("design") or {}
    lines = []
    kind = design.get("kind", "sweep")
    axes = design.get("axes") or {}
    if axes:
        lines.append("Design: " + kind + " over " + "; ".join(
            f"{k} ∈ {{{', '.join(str(x) for x in v)}}}" for k, v in axes.items()))
    elif design.get("conditions"):
        lines.append(f"Design: {kind} · {len(design['conditions'])} named conditions")
    else:
        lines.append(f"Design: {kind}")
    fixed = exp.get("fixed") or {}
    if fixed:
        lines.append("Fixed parameters: " + ", ".join(f"{k} = {v}" for k, v in fixed.items()))
    seeds = exp.get("seeds", design.get("seeds"))
    if seeds:
        lines.append(f"Replications: {seeds} seeds per cell")
    metric = exp.get("metric") or {}
    if metric.get("name"):
        reduce = str(metric.get("reduce", "")).strip()
        lines.append(f"Primary metric — {metric['name']}" + (f": {reduce}" if reduce else ""))
    for sm in (metric.get("secondary") or []):
        if sm.get("name"):
            r = str(sm.get("reduce", "")).strip()
            lines.append(f"Secondary — {sm['name']}" + (f": {r}" if r else ""))
    return lines


def _unrendered_outputs(exp: dict) -> list:
    """Preregistered outputs rayleigh can't draw itself (artifact / stat / …), so the report
    names them rather than letting a planned deliverable silently vanish."""
    out = []
    for o in exp.get("outputs") or []:
        if o.get("kind") in ("figure", "table"):
            continue
        out.append((o.get("kind") or "output", o.get("name") or "",
                    str(o.get("caption") or "").strip()))
    return out


def _code_sha_from_provenance(results: Path) -> str | None:
    """Best-effort code SHA for the provenance footer: read the first cell's .prov.json."""
    data = results / "data"
    if not data.is_dir():
        return None
    for p in data.rglob("*.prov.json"):
        try:
            return json.loads(p.read_text()).get("code_sha")
        except Exception:                                # noqa: BLE001
            return None
    return None


def _build_docx(path: Path, project, cycle, brief, items, author,
                generated: datetime, code_sha: str | None) -> bool:
    """The human deliverable: a full, self-contained report — title page, an executive summary
    across experiments, then one complete section per experiment (question, method, the
    preregistration, results prose + metric ranges, embedded numbered figures/tables, and any
    planned artifacts rayleigh couldn't draw), closing with provenance."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError:
        return False

    def italic(text, size=None):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.italic = True
        if size:
            r.font.size = Pt(size)
        return p

    def bullet(text):
        doc.add_paragraph(text, style="List Bullet")

    doc = Document()
    doc.core_properties.author = author

    # ── Title page ──────────────────────────────────────────────────────────────────
    doc.add_heading(f"{project}", 0)
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run(f"Results — research cycle {cycle}")
    r.bold = True
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Generated {generated:%Y-%m-%d}").italic = True
    if brief:
        doc.add_heading("Overview", level=1)
        doc.add_paragraph(brief.strip())

    # ── Executive summary across experiments ────────────────────────────────────────
    if len(items) > 1:
        doc.add_heading("Summary of experiments", level=1)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        for j, h in enumerate(("ID", "Experiment", "Mode", "Coverage")):
            run = t.rows[0].cells[j].paragraphs[0].add_run(h)
            run.bold = True
        for it in items:
            exp = it["exp"]
            cells = t.add_row().cells
            cells[0].text = str(exp["id"])
            cells[1].text = str(exp.get("title", ""))
            cells[2].text = str(exp.get("mode", "confirmatory"))
            cells[3].text = f"{it['n_data']}/{it['n_cells']} cells"

    # ── One full section per experiment ─────────────────────────────────────────────
    fig_no = tab_no = 0
    for it in items:
        exp = it["exp"]
        doc.add_page_break()
        doc.add_heading(f"{exp['id']} — {exp.get('title', '')}", level=1)
        italic(f"{str(exp.get('mode', 'confirmatory')).capitalize()} · "
               f"{it['n_data']}/{it['n_cells']} cells with data")

        if exp.get("question"):
            doc.add_heading("Question", level=2)
            doc.add_paragraph(str(exp["question"]).strip())

        method = _methods_lines(exp)
        if method:
            doc.add_heading("Method", level=2)
            for line in method:
                bullet(line)

        expected = exp.get("expected_direction")
        if str(exp.get("mode", "confirmatory")) == "confirmatory" and expected:
            doc.add_heading("Preregistered expectation", level=2)
            doc.add_paragraph(str(expected).strip())
        elif str(exp.get("mode")) == "exploratory":
            italic("Exploratory experiment — not preregistered; findings are hypothesis-generating.")

        doc.add_heading("Results", level=2)
        f = doc.add_paragraph()
        f.add_run("Finding: ").bold = True
        f.add_run(it["finding"])
        an = it.get("analysis") or {}
        if an.get("script") is not None:
            italic(f"Analytical products generated by results/analysis/{Path(an['script']).name} "
                   f"(base R + ggplot2) — re-run `Rscript results/analysis/{Path(an['script']).name}` "
                   f"to regenerate.")

        for primary, _all, cap in it["figures"]:
            fig_no += 1
            try:
                doc.add_picture(str(primary), width=Inches(6.0))   # PNG — docx can't embed SVG
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:                       # noqa: BLE001
                doc.add_paragraph(f"[figure {fig_no} could not be embedded: {e}]")
            cap_p = doc.add_paragraph()
            cr = cap_p.add_run(f"Figure {fig_no}. {cap}")
            cr.italic = True
            cr.font.size = Pt(9)

        for k, (piv, cap) in enumerate(it["tables"]):
            tab_no += 1
            t = doc.add_table(rows=1, cols=len(piv.columns) + 1)
            t.style = "Table Grid"
            hdr = t.rows[0].cells
            hdr[0].paragraphs[0].add_run(str(piv.index.name or "")).bold = True
            for j, c in enumerate(piv.columns):
                hdr[j + 1].paragraphs[0].add_run(str(c)).bold = True
            for idx, row in piv.iterrows():
                cells = t.add_row().cells
                cells[0].text = str(idx)
                for j, v in enumerate(row):
                    cells[j + 1].text = _fmt(v)
            cap_p = doc.add_paragraph()
            cr = cap_p.add_run(f"Table {tab_no}. {cap}")
            cr.italic = True
            cr.font.size = Pt(9)

        planned = _unrendered_outputs(exp)
        if planned:
            doc.add_heading("Planned outputs (produced outside this report)", level=2)
            for kind, name, cap in planned:
                label = f"{kind}" + (f" — {name}" if name else "")
                bullet(f"{label}: {cap}" if cap else label)

    # ── Synthesis (collated, not adjudicated) ───────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Synthesis", level=1)
    n = len(items)
    conf = sum(1 for it in items if str(it["exp"].get("mode", "confirmatory")) == "confirmatory")
    tot_data = sum(it["n_data"] for it in items)
    tot_cells = sum(it["n_cells"] for it in items)
    doc.add_paragraph(
        f"This cycle comprises {n} experiment{'s' if n != 1 else ''} "
        f"({conf} confirmatory, {n - conf} exploratory), {tot_data}/{tot_cells} cells with data. "
        f"Each finding below is stated against its own preregistration; the findings are brought "
        f"together here but not adjudicated across experiments.")
    doc.add_heading("Findings at a glance", level=2)
    for it in items:
        exp = it["exp"]
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{exp['id']} — {exp.get('title', '')}: ").bold = True
        p.add_run(it["finding"])
    deps = [(str(it["exp"]["id"]), it["exp"].get("depends_on"))
            for it in items if it["exp"].get("depends_on")]
    if deps:
        doc.add_paragraph("Design dependencies: "
                          + "; ".join(f"{a} builds on {b}" for a, b in deps) + ".")
    italic("Cross-experiment interpretation — reconciling these findings into a single account — "
           "belongs to the author's write-up (or raconteur); rayleigh reports, it does not conclude.")

    # ── Provenance ──────────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Provenance", level=1)
    bullet(f"Generated by rayleigh {__version__} on {generated:%Y-%m-%d %H:%M UTC}.")
    if code_sha and code_sha != "unknown":
        bullet(f"Model/analysis code at commit {code_sha}.")
    scripts = [Path(it["analysis"]["script"]).name for it in items
               if (it.get("analysis") or {}).get("script")]
    if scripts:
        bullet("Every figure, table, and statistic was generated by base R + ggplot2, not by "
               "rayleigh — regenerate any time with: "
               + "; ".join(f"Rscript results/analysis/{s}" for s in sorted(set(scripts))) + ".")
    bullet("Tidy per-cell data in results/analysis/data/*.csv; figures as PNG + EPS in "
           "results/figures/; table/stat CSVs in results/tables/ and results/analysis/stats/; "
           "structured findings in results/findings.json.")
    doc.save(str(path))
    return True


# --------------------------------------------------------------------- command
def run_process_outputs(args) -> int:
    root = Path(args.dir).resolve() if getattr(args, "dir", None) else Path.cwd()
    results = root / "results"
    spec_path = active_spec_path(results / "designdocs")
    if not spec_path.is_file():
        log(f"no {spec_path} — run `rayleigh init` first")
        return 1
    spec = yaml.safe_load(spec_path.read_text()) or {}
    exps = spec.get("experiments") or []
    if not exps:
        log("experiments.yaml has no experiments — author them in `rayleigh init`")
        return 1
    if getattr(args, "experiment", None):
        exps = [e for e in exps if str(e.get("id")) == args.experiment]
        if not exps:
            log(f"experiment '{args.experiment}' not found")
            return 1

    meta = {}
    ry = results / "rayleigh.yaml"
    if ry.is_file():
        meta = yaml.safe_load(ry.read_text()) or {}
    code = spec.get("code") or {}
    adapter = code.get("run_adapter") or {}
    code_dir = (results / (code.get("path") or "../code")).resolve()
    cycle = str(spec.get("cycle") or meta.get("cycle") or date.today().strftime("%y%m%d"))
    project = str(spec.get("project") or meta.get("project") or "project")
    brief = str(spec.get("brief") or meta.get("brief") or "")

    from rayleigh.config import load_config
    cfg = load_config()
    loader = _get_loader(spec, code_dir, results)   # figures/tables dirs are created by r_analysis

    dry = getattr(args, "dry_run", False)
    items = []
    for exp in exps:
        eid = str(exp["id"])
        df, param_cols, metric_cols, missing = build_table(exp, adapter, results, loader)
        n_cells = len(df) + missing
        planned = [o.get("kind") for o in exp.get("outputs") or []]
        if dry:
            log(f"{eid}: {len(df)}/{n_cells} cells have data · outputs planned: {planned}")
            continue
        if df.empty:
            log(f"{eid}: no cell data yet — run `rayleigh conduct_exp {eid}` first")
            items.append({"exp": exp, "finding": "No data collected yet.", "figures": [],
                          "tables": [], "table_files": [], "n_data": 0, "n_cells": n_cells,
                          "analysis": None})
            continue

        # rayleigh doesn't render the analytical products — it writes the tidy data + a base-R
        # script and runs it. The report is assembled from whatever R produced.
        res = r_analyze(eid, exp, df, param_cols, metric_cols, results, custom_r=code.get("analysis_r"))
        if not res["ran"]:
            log(f"  {eid}: {res['log']} — wrote data + results/analysis/{eid}.R; install R + "
                f"ggplot2 and re-run `Rscript results/analysis/{eid}.R` to render.")
        elif not res["ok"]:
            log(f"  {eid}: R analysis reported an error (outputs may be partial):\n{res['log']}")
        for note, name, cap in res["unrendered"]:
            log(f"  {eid}: output '{note}'{f' ({name})' if name else ''} isn't derivable from cell "
                f"data — name it in the write-up (or add a `regression:`/adapter helper).")
        log(f"{eid}: {len(df)}/{n_cells} cells · {len(res['figures'])} figure(s), "
            f"{len(res['tables'])} table(s) via R{'' if res['ok'] else ' [R incomplete]'}")
        items.append({"exp": exp, "finding": finding_text(exp, res["summary_df"]),
                      "figures": res["figures"], "tables": res["tables"],
                      "table_files": res["table_files"], "n_data": len(df), "n_cells": n_cells,
                      "analysis": {"script": res["script"], "data": res["data"],
                                   "ran": res["ran"], "ok": res["ok"]}})
    if dry:
        return 0

    (results / "RESULTS.md").write_text(_build_markdown(project, cycle, brief, items, results))
    log(f"wrote {(results / 'RESULTS.md').relative_to(root)}")

    # Structured hand-off for raconteur (its load_results ingests results/*.json + *.csv).
    (results / "findings.json").write_text(_build_findings(project, cycle, brief, items, results))
    log(f"wrote {(results / 'findings.json').relative_to(root)}"
        + (f" + {len([f for it in items for f in it.get('table_files', [])])} table CSV(s)"
           if any(it.get("table_files") for it in items) else ""))

    if getattr(args, "no_docx", False):
        log("--no-docx: skipped the .docx (RESULTS.md written)")
        log("next: `rayleigh review` — the human-led review of this report before the cycle closes.")
        return 0
    docx_name = f"{cycle}_{project}_results_{cfg.tool_initials}.docx"
    if _build_docx(results / docx_name, project, cycle, brief, items, cfg.author_name,
                   datetime.now(timezone.utc), _code_sha_from_provenance(results)):
        log(f"wrote results/{docx_name}")
    else:
        log("python-docx unavailable — wrote RESULTS.md only "
            "(pip install python-docx for the .docx deliverable)")
    log("next: `rayleigh review` — the human-led review of this report before the cycle closes.")
    return 0
